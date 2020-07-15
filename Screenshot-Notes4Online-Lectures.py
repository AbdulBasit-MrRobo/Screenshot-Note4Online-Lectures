#by AbdulBasit Hakimi
import docx
from datetime import date
import keyboard
from PIL import ImageGrab
import os  
mydoc = docx.Document()
while True:

    if keyboard.is_pressed('n') :

        im = ImageGrab.grab(bbox = None)
        im.save("write the path for image here")
        mydoc.add_picture("write the path of above image",width=docx.shared.Inches(6.26), height=docx.shared.Inches(3.386))
        mydoc.add_paragraph("")
        mydoc.add_paragraph("")
        os.remove("write the path of above image")
    
    elif keyboard.is_pressed('s') :

        mydoc.save("write the path for doc file" + str(date.today()) +".docx")

    elif keyboard.is_pressed('q') :
        quit()
